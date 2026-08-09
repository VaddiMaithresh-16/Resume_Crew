"""Safe document-to-text conversion for supported application inputs."""

from __future__ import annotations

from pathlib import Path

MAX_INPUT_FILE_SIZE_MB = 15
MAX_EXTRACTED_CHARACTERS = 100_000
SUPPORTED_DOCUMENT_EXTENSIONS = (".pdf", ".docx", ".txt", ".md")


def _validate_file(path: Path) -> None:
    if not path.exists():
        raise FileNotFoundError(f"File not found: {path}")
    if not path.is_file():
        raise ValueError(f"'{path}' is not a file.")
    size_mb = path.stat().st_size / (1024 * 1024)
    if size_mb > MAX_INPUT_FILE_SIZE_MB:
        raise ValueError(
            f"'{path}' is {size_mb:.1f} MB; the limit is {MAX_INPUT_FILE_SIZE_MB} MB."
        )


def _validate_text(text: str, path: Path) -> str:
    cleaned = text.strip()
    if not cleaned:
        raise ValueError(f"No extractable text found in '{path}'. It may require OCR.")
    if len(cleaned) > MAX_EXTRACTED_CHARACTERS:
        raise ValueError(
            f"'{path}' contains {len(cleaned):,} characters after extraction, exceeding "
            f"the {MAX_EXTRACTED_CHARACTERS:,}-character analysis limit. Split or shorten it."
        )
    return cleaned


def _read_pdf(path: Path) -> str:
    from pypdf import PdfReader
    from pypdf.errors import PdfReadError

    try:
        reader = PdfReader(str(path))
        if reader.is_encrypted and reader.decrypt("") == 0:
            raise ValueError(f"'{path}' is password-protected. Remove the password and try again.")
        return _validate_text("\n".join(page.extract_text() or "" for page in reader.pages), path)
    except ValueError:
        raise
    except PdfReadError as exc:
        raise ValueError(f"'{path}' is not a valid PDF: {exc}") from exc
    except Exception as exc:
        raise ValueError(f"Could not extract text from '{path}': {exc}") from exc


def _read_docx(path: Path) -> str:
    import docx
    from docx.opc.exceptions import PackageNotFoundError

    try:
        document = docx.Document(str(path))
    except PackageNotFoundError as exc:
        raise ValueError(f"'{path}' is not a valid .docx file: {exc}") from exc

    parts: list[str] = []

    # Extract paragraphs with safe encoding handling.
    for paragraph in document.paragraphs:
        text = paragraph.text
        if not isinstance(text, str):
            continue
        # Replace any problematic characters rather than crashing.
        text = text.encode("utf-8", errors="replace").decode("utf-8", errors="replace")
        if text.strip():
            parts.append(text)

    # Extract tables with structural context so the LLM understands the data.
    for table in document.tables:
        rows = table.rows
        if not rows:
            continue

        table_lines: list[str] = []
        for row_idx, row in enumerate(rows):
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if not cells:
                continue
            line = " | ".join(cells)
            # Treat the first row as a header and mark it clearly.
            if row_idx == 0:
                table_lines.append(f"[TABLE HEADER] {line}")
            else:
                table_lines.append(f"[TABLE ROW] {line}")

        if table_lines:
            parts.extend(table_lines)

    return _validate_text("\n".join(part for part in parts if part.strip()), path)


def extract_text(file_path: str) -> str:
    """Extract validated text from a PDF, DOCX, TXT, or Markdown file."""
    path = Path(file_path).expanduser().resolve()
    _validate_file(path)
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _read_pdf(path)
    if suffix == ".docx":
        return _read_docx(path)
    if suffix in (".txt", ".md"):
        return _validate_text(path.read_text(encoding="utf-8", errors="replace"), path)
    raise ValueError(
        f"Unsupported file type '{suffix}'. Supported types: {', '.join(SUPPORTED_DOCUMENT_EXTENSIONS)}"
    )
